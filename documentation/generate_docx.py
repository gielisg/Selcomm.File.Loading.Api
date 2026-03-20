"""Generate architecture.docx from the architecture content."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# -- Style tweaks --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

code_style = doc.styles.add_style('CodeBlock', 1)  # paragraph style
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.left_indent = Cm(0.5)

def add_code(text):
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='CodeBlock')

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Selcomm.File.Loading.Api')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Architecture Document')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Module: ').bold = True
meta.add_run('file-loading   ')
meta.add_run('Port: ').bold = True
meta.add_run('5140   ')
meta.add_run('Framework: ').bold = True
meta.add_run('.NET 10.0   ')
meta.add_run('Database: ').bold = True
meta.add_run('Informix (ODBC)')

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (placeholder)
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. High-Level Architecture',
    '3. Project Structure',
    '4. API Surface',
    '5. Core Subsystems',
    '   5.1 File Parsing Engine',
    '   5.2 Validation Engine',
    '   5.3 File Transfer System',
    '   5.4 Background Worker',
    '   5.5 Activity Audit System',
    '   5.6 Generic Configurable Parser',
    '6. Data Model',
    '7. Authentication & Security',
    '8. Configuration',
    '9. Dependencies',
    '10. Deployment',
    '11. Design Patterns',
    '12. Key Architectural Decisions',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. OVERVIEW
# ============================================================
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'The File Loading API is a V4 REST API that modernises the legacy ntfileload.4gl system. '
    'It converts a batch-oriented 4GL program into an event-driven, memory-efficient REST service '
    'with streaming two-pass file processing.'
)
doc.add_paragraph('The API provides:')
bullets = [
    ('File loading', 'Parse and load CDR, CHG, EBL, SVC, and ORD files into the Selcomm database'),
    ('File transfer management', 'Automated and on-demand file downloads from SFTP/FTP/FileSystem sources'),
    ('Workflow management', 'Track files through Transfer > Processing > Processed/Errors/Skipped folders'),
    ('Validation engine', 'Configurable file-level and field-level validation with AI-friendly error reporting'),
    ('Activity auditing', 'Full audit trail of all file operations'),
]
for title, desc in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title} — ').bold = True
    p.add_run(desc)

# ============================================================
# 2. HIGH-LEVEL ARCHITECTURE
# ============================================================
doc.add_heading('2. High-Level Architecture', level=1)
doc.add_paragraph(
    'The following diagram shows the layered architecture of the File Loading API, '
    'from the Swagger UI at the top through the ASP.NET Core pipeline, controllers, '
    'services, shared infrastructure, repository, and database.'
)

arch_diagram = """\
                         Swagger UI (/swagger)
                                |
               ASP.NET Core Pipeline (Auth, CORS, Routing)
                                |
          +---------------------+---------------------+
          |                                           |
  FileLoaderController                FileManagementController
  /api/v4/fileloader                  /api/v4/filemanager
          |                                           |
  IFileLoaderService            IFileManagementService
                                IFileTransferService
          |                                           |
          +---------------------+---------------------+
                                |
                    Shared Infrastructure
          +-----------+-----------+-----------+
          |           |           |           |
     File Parsers  Validation  Transfer   Background
     (Strategy)    Engine      Clients    Worker
          |           |           |
     Error         Compression
     Aggregator    Helper
                                |
                   IFileLoaderRepository (Dapper/ODBC)
                                |
                      Informix Database
             (nt_file, cl_detail, ntfl_transfer, ...)\
"""
add_code(arch_diagram)

# ============================================================
# 3. PROJECT STRUCTURE
# ============================================================
doc.add_heading('3. Project Structure', level=1)

structure = """\
FileLoading/
 +-- Controllers/
 |   +-- FileLoaderController.cs        Core file loading (6 endpoints)
 |   +-- FileManagementController.cs    Transfer/workflow (22 endpoints)
 +-- Interfaces/
 |   +-- IFileLoaderService.cs          File loading contract
 |   +-- IFileManagementService.cs      Management contract
 |   +-- IFileTransferService.cs        Transfer contract
 +-- Services/
 |   +-- FileLoaderService.cs           Parsing, validation, DB insertion
 |   +-- FileManagementService.cs       Workflow and dashboard
 |   +-- FileTransferService.cs         Remote downloads and config
 +-- Repositories/
 |   +-- IFileLoaderRepository.cs       Data access contract (50+ methods)
 |   +-- FileLoaderRepository.cs        ODBC/Dapper implementation
 +-- Models/
 |   +-- FileLoaderModels.cs            Requests, responses, configuration
 |   +-- StagingModels.cs               Database records (cl_detail, etc.)
 |   +-- GenericParserModels.cs        Generic parser config and models
 |   +-- TransferModels.cs              Transfer, folder, activity models
 |   +-- ValidationModels.cs            Validation rules, results, AI summaries
 +-- Parsers/
 |   +-- BaseFileParser.cs              Abstract template method base
 |   +-- ChgFileParser.cs               Charge file parser
 |   +-- EblFileParser.cs               Equipment/billing parser
 |   +-- OrdFileParser.cs               Order file parser
 |   +-- SvcFileParser.cs               Service record parser
 |   +-- GenericFileParser.cs          Data-driven configurable parser
 |   +-- FileRowReaders.cs            Row reader abstraction (text/Excel)
 |   +-- Cdr/
 |       +-- GenericCdrParser.cs        Selcomm pipe-delimited CDR
 |       +-- TelstraGsmCdrParser.cs     Telstra GSM
 |       +-- TelstraCdmaCdrParser.cs    Telstra CDMA
 |       +-- OptusCdrParser.cs          Optus
 |       +-- AaptCdrParser.cs           AAPT
 |       +-- VodafoneCdrParser.cs       Vodafone
 +-- Transfer/
 |   +-- ITransferClient.cs            Protocol-agnostic interface
 |   +-- TransferClientFactory.cs      Factory for protocol selection
 |   +-- SftpTransferClient.cs         SSH.NET SFTP
 |   +-- FtpTransferClient.cs          FluentFTP
 |   +-- FileSystemTransferClient.cs   Local/network paths
 |   +-- CompressionHelper.cs          GZip and Zip
 +-- Validation/
 |   +-- ValidationEngine.cs           Field-level validation
 |   +-- IValidationConfigProvider.cs  Config source interface
 |   +-- ValidationConfigProvider.cs   Config from appsettings/DB
 |   +-- ErrorAggregator.cs            Smart error summarisation
 +-- Workers/
 |   +-- FileTransferWorker.cs         Scheduled download service
 +-- Data/
 |   +-- FileLoaderDbContext.cs        ODBC database context
 +-- Program.cs                        Startup and DI
 +-- appsettings.json                  Module configuration\
"""
add_code(structure)

# ============================================================
# 4. API SURFACE
# ============================================================
doc.add_heading('4. API Surface', level=1)

doc.add_heading('4.1 FileLoaderController (/api/v4/fileloader)', level=2)
doc.add_paragraph('Core file loading endpoints — the modernised replacement for ntfileload.4gl.')
add_table(
    ['Method', 'Route', 'Purpose'],
    [
        ['POST', '/load', 'Load a network file by path'],
        ['POST', '/upload', 'Upload a file via multipart form and load it'],
        ['GET', '/files/{nt-file-num}', 'Get load status for a file'],
        ['GET', '/files', 'List loaded files with filtering'],
        ['GET', '/file-types', 'List supported file type codes'],
        ['POST', '/files/{nt-file-num}/reprocess', 'Reprocess a previously loaded file'],
    ],
    col_widths=[2, 5.5, 7]
)

doc.add_heading('4.2 FileManagementController (/api/v4/filemanager)', level=2)
doc.add_paragraph('File transfer management, workflow operations, and operational monitoring.')

doc.add_heading('Dashboard & Files', level=3)
add_table(
    ['Method', 'Route', 'Purpose'],
    [
        ['GET', '/dashboard', 'Summary counts and status overview'],
        ['GET', '/files', 'List files with filtering and pagination'],
        ['GET', '/files/{transfer-id}', 'Get file details by transfer ID'],
        ['POST', '/files/{transfer-id}/process', 'Trigger processing of a transferred file'],
        ['POST', '/files/{transfer-id}/retry', 'Retry a failed file'],
        ['POST', '/files/{transfer-id}/move', 'Move file to a workflow folder'],
        ['POST', '/files/{nt-file-num}/unload', 'Reverse a file load (delete inserted records)'],
        ['POST', '/files/{nt-file-num}/skip-sequence', 'Force skip to a sequence number'],
        ['GET', '/files/{transfer-id}/download', 'Download file content to browser'],
        ['DELETE', '/files/{transfer-id}', 'Delete a file record'],
    ],
    col_widths=[2, 5.5, 7]
)

doc.add_heading('Activity, Validation & Exceptions', level=3)
add_table(
    ['Method', 'Route', 'Purpose'],
    [
        ['GET', '/activity', 'Query activity audit log'],
        ['GET', '/files/{nt-file-num}/validation-summary', 'AI-friendly validation error summary'],
        ['GET', '/exceptions/errors', 'List files with processing errors'],
        ['GET', '/exceptions/skipped', 'List manually skipped files'],
    ],
    col_widths=[2, 6.5, 6]
)

doc.add_heading('Transfer Sources & Folders', level=3)
add_table(
    ['Method', 'Route', 'Purpose'],
    [
        ['POST', '/transfers/{source-id}/fetch', 'Fetch files from a transfer source'],
        ['GET', '/sources', 'List configured transfer sources'],
        ['GET', '/sources/{source-id}', 'Get source configuration'],
        ['PUT', '/sources/{source-id}', 'Create or update a source'],
        ['DELETE', '/sources/{source-id}', 'Delete a source'],
        ['POST', '/sources/{source-id}/test', 'Test connection to a saved source'],
        ['POST', '/sources/test', 'Test connection with ad-hoc config'],
        ['GET', '/folders', 'Get folder workflow configuration'],
        ['PUT', '/folders', 'Save folder workflow configuration'],
    ],
    col_widths=[2, 5.5, 7]
)

doc.add_heading('Parser Configuration', level=3)
add_table(
    ['Method', 'Route', 'Purpose'],
    [
        ['GET', '/parsers', 'List all generic parser configs'],
        ['GET', '/parsers/{file-type-code}', 'Get a parser config with column mappings'],
        ['PUT', '/parsers/{file-type-code}', 'Create or update a parser config'],
        ['DELETE', '/parsers/{file-type-code}', 'Delete a parser config'],
    ],
    col_widths=[2, 5.5, 7]
)

# ============================================================
# 5. CORE SUBSYSTEMS
# ============================================================
doc.add_heading('5. Core Subsystems', level=1)

# 5.1 File Parsing
doc.add_heading('5.1 File Parsing Engine', level=2)
doc.add_paragraph('Pattern: Strategy + Template Method')
doc.add_paragraph(
    'The parsing subsystem uses an abstract BaseFileParser that defines the two-pass streaming '
    'workflow, with concrete implementations for each file format.'
)

add_table(
    ['Parser', 'File Type', 'Format'],
    [
        ['GenericCdrParser', 'CDR', 'Selcomm pipe-delimited (H|D|T)'],
        ['TelstraGsmCdrParser', 'TEL_GSM', 'Telstra GSM CDR'],
        ['TelstraCdmaCdrParser', 'TEL_CDMA', 'Telstra CDMA CDR'],
        ['OptusCdrParser', 'OPTUS', 'Optus CDR'],
        ['AaptCdrParser', 'AAPT', 'AAPT CDR'],
        ['VodafoneCdrParser', 'VODA', 'Vodafone CDR'],
        ['ChgFileParser', 'CHG', 'Charge detail records'],
        ['SvcFileParser', 'SVC', 'Service records'],
        ['OrdFileParser', 'ORD', 'Order records'],
        ['EblFileParser', 'EBL', 'Equipment/billing records'],
        ['GenericFileParser', 'GEN', 'Data-driven configurable (CSV/XLSX/delimited)'],
    ],
    col_widths=[4.5, 3, 7]
)

doc.add_heading('Two-Pass Streaming Processing', level=3)
doc.add_paragraph(
    'Pass 1 (Validate): Stream the file and validate structure (header, trailer, sequence, '
    'field constraints) without loading all records into memory.'
)
doc.add_paragraph(
    'Pass 2 (Insert): Re-read the file and insert valid records in configurable batches '
    'with explicit transaction batching.'
)
doc.add_paragraph(
    'This approach reduces memory usage from ~280MB to ~10-50MB for a 400K-record file.'
)

doc.add_heading('File Record Types', level=3)
add_table(
    ['Type', 'Code', 'Description'],
    [
        ['Header', 'H', 'File metadata, sequence numbers'],
        ['Detail', 'D', 'Data records (CDRs, charges, etc.)'],
        ['Trailer', 'T', 'Record counts, cost totals for reconciliation'],
    ],
    col_widths=[3, 2, 9.5]
)

# 5.2 Validation Engine
doc.add_heading('5.2 Validation Engine', level=2)

add_table(
    ['Component', 'Responsibility'],
    [
        ['ValidationEngine', 'Executes field-level validation rules against parsed records'],
        ['ValidationConfigProvider', 'Loads validation configuration from appsettings or database'],
        ['ErrorAggregator', 'Collects errors, aggregates after threshold, produces AI summaries'],
    ],
    col_widths=[4.5, 10]
)

doc.add_heading('Validation Levels', level=3)
doc.add_paragraph('File-level:', style='List Bullet').runs[0].bold = True
doc.add_paragraph(
    'Header/trailer presence, sequence contiguity, footer count matching, min/max record counts',
    style='List Bullet 2' if 'List Bullet 2' in [s.name for s in doc.styles] else 'List Bullet'
)
doc.add_paragraph('Field-level:', style='List Bullet').runs[0].bold = True
doc.add_paragraph(
    'Type parsing (string, int, long, decimal, datetime, bool), required checks, range constraints, '
    'regex patterns, allowed value sets, date boundary checks',
    style='List Bullet'
)

doc.add_heading('Error Aggregation Strategy', level=3)
agg_items = [
    'First 100 errors are stored with full detail (line number, raw data, field name)',
    'Beyond threshold, errors are aggregated by (ErrorCode, FieldName) with sample line numbers',
    'File-level errors always retain full detail',
    'Raw data is truncated to 500 characters',
]
for item in agg_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('AI-Friendly Output (ValidationSummaryForAI)', level=3)
doc.add_paragraph(
    'The validation engine generates plain-English summaries designed for consumption by AI agents '
    'and LLM-powered support tools:'
)
add_code("""\
{
  "OverallStatus": "File has 47 errors across 3 fields",
  "MainIssues": [
    "23 records have invalid date format in CallStartDate",
    "18 records have negative values in Cost field"
  ],
  "ErrorCountsByField": { "CallStartDate": 23, "Cost": 18, "Duration": 6 },
  "SuggestedActions": ["Fix date format in source system"],
  "CanPartiallyProcess": true
}\
""")

# 5.3 File Transfer System
doc.add_heading('5.3 File Transfer System', level=2)
doc.add_paragraph('Pattern: Factory + Strategy')

add_table(
    ['Client', 'Protocol', 'Library'],
    [
        ['SftpTransferClient', 'SFTP', 'SSH.NET'],
        ['FtpTransferClient', 'FTP', 'FluentFTP'],
        ['FileSystemTransferClient', 'FileSystem', 'System.IO'],
    ],
    col_widths=[5, 3, 6.5]
)

doc.add_heading('Transfer Workflow', level=3)
add_code("""\
Remote Source --fetch--> Transfer Folder --process--> Processing Folder
                                                          |
                                              +-----------+-----------+
                                              v           v           v
                                          Processed    Errors      Skipped\
""")

doc.add_paragraph('Capabilities:', style='List Bullet').runs[0].bold = True
capabilities = [
    'Protocol-agnostic file listing with glob pattern matching',
    'Skip-pattern support (ignore files matching pattern)',
    'Duplicate detection via ntfl_downloaded_files table',
    'Optional compression on archive (GZip or Zip)',
    'Connection testing (saved and ad-hoc configurations)',
]
for cap in capabilities:
    doc.add_paragraph(cap, style='List Bullet')

# 5.4 Background Worker
doc.add_heading('5.4 Background Worker (FileTransferWorker)', level=2)
doc.add_paragraph(
    'A BackgroundService that runs on a 1-minute tick. It queries all enabled transfer sources, '
    'evaluates each source\'s CRON schedule (via the Cronos library), and executes '
    'FetchFilesFromSourceAsync when a schedule fires. Creates a system-level security context '
    'for automated operations. Recovers from errors without stopping the worker.'
)
doc.add_paragraph('CRON Format: Standard 5-field (minute, hour, day-of-month, month, day-of-week)')

# 5.5 Activity Audit
doc.add_heading('5.5 Activity Audit System', level=2)
doc.add_paragraph(
    'Every significant operation is logged to ntfl_activity_log with the activity type, '
    'associated file/transfer IDs, user ID, domain, and a JSON details payload.'
)

add_table(
    ['ID', 'Activity Type', 'ID', 'Activity Type'],
    [
        ['1', 'Downloaded', '9', 'FileDeleted'],
        ['2', 'MovedToProcessing', '10', 'FileUnloaded'],
        ['3', 'ProcessingStarted', '11', 'SequenceSkipped'],
        ['4', 'ProcessingCompleted', '12', 'ManualDownload'],
        ['5', 'ProcessingFailed', '13', 'BrowserDownload'],
        ['6', 'MovedToSkipped', '14', 'SourceCreated'],
        ['7', 'MovedToErrors', '15', 'SourceModified'],
        ['8', 'MovedToProcessed', '16', 'SourceDeleted'],
    ],
    col_widths=[1.5, 5, 1.5, 5]
)

# 5.6 Generic Configurable Parser
doc.add_heading('5.6 Generic Configurable Parser', level=2)
doc.add_paragraph(
    'Supports many small vendors/networks that send files as CSV, Excel, or text without '
    'requiring a new parser class and DB table per vendor. Configuration is entirely data-driven '
    '- add a vendor by inserting database config rows, not by writing code.'
)

doc.add_heading('Database-Driven Configuration', level=3)
doc.add_paragraph(
    'Two configuration tables drive the parser:'
)
config_tables = [
    ('ntfl_file_format_config', 'One row per file type: format (CSV/XLSX/Delimited), delimiter, header row flag, skip rows top/bottom, row identification mode, trailer total configuration, optional custom SP'),
    ('ntfl_column_mapping', 'One row per column per file type: source column index to target field name, data type, validation (required, regex, max length), date format override, default value'),
]
for name, desc in config_tables:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{name}: ').bold = True
    p.add_run(desc)

doc.add_heading('Row Identification Modes', level=3)
add_table(
    ['Mode', 'How it works', 'Best for'],
    [
        ['Position', 'Skip first N rows, skip last N. First non-skipped row is header if has_header_row=Y. Trailer detected by indicator pattern.', 'Simple CSV exports with fixed structure'],
        ['Indicator', 'Read value in column row_id_column. Compare against header/trailer/skip/detail indicator strings.', 'Files with a record type column (H/D/T format)'],
        ['Pattern', 'Apply regex from indicator fields against the full raw line.', 'Complex files where row type is determined by content patterns'],
    ],
    col_widths=[2.5, 7, 5]
)

doc.add_heading('Standard Target Fields', level=3)
doc.add_paragraph(
    'AccountCode, ServiceId, ChargeType, CostAmount, TaxAmount, Quantity, UOM, FromDate, '
    'ToDate, Description, ExternalRef - plus Generic01..Generic20 overflow columns for '
    'vendor-specific fields. All records are stored in a single ntfl_generic_detail staging table.'
)

doc.add_heading('Row Reader Abstraction', level=3)
doc.add_paragraph(
    'The IFileRowReader interface (Strategy pattern) decouples file format from parsing logic:'
)
reader_items = [
    ('DelimitedTextRowReader', 'Reads CSV, pipe-delimited, tab-delimited, and semicolon-delimited text files'),
    ('ExcelRowReader', 'Reads .xlsx files using ClosedXML (MIT licensed). Supports sheet selection by name or index.'),
]
for name, desc in reader_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{name}: ').bold = True
    p.add_run(desc)

doc.add_heading('Additional Features', level=3)
features = [
    'Total reconciliation: supports both SUM (cost total) and COUNT (record count) against a configurable trailer column',
    'Custom SP hook: after generic records are inserted, an optional stored procedure can be called for vendor-specific validation',
    'Automatic parser fallback: when no dedicated parser matches a file type, the service checks for generic config in the database',
]
for item in features:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 6. DATA MODEL
# ============================================================
doc.add_heading('6. Data Model', level=1)

doc.add_heading('6.1 Legacy Tables (from 4GL system)', level=2)
add_table(
    ['Table', 'Purpose'],
    [
        ['nt_file', 'Master file records (file number, type, status, dates)'],
        ['nt_file_stat', 'File status lookup values'],
        ['file_type', 'File type definitions'],
        ['cl_detail', 'Call detail records (CDR data)'],
        ['nt_fl_trailer', 'File trailer totals (reconciliation)'],
        ['nt_cl_not_load', 'Records that failed to load (with error reason)'],
    ],
    col_widths=[4, 10.5]
)

doc.add_heading('6.2 V4 Tables (new for this module)', level=2)
add_table(
    ['Table', 'Purpose'],
    [
        ['ntfl_transfer_source', 'Transfer source configurations (SFTP/FTP/FS)'],
        ['ntfl_transfer', 'File transfer tracking (status, folder, timestamps)'],
        ['ntfl_folder_config', 'Folder workflow paths per domain/file-type'],
        ['ntfl_downloaded_files', 'Downloaded file cache (prevents re-downloads)'],
        ['ntfl_activity_log', 'Audit trail for all operations'],
        ['ntfl_validation_summary', 'AI-friendly validation results (JSON)'],
        ['ntfl_error_log', 'Detailed parse/validation errors'],
        ['ntfl_file_format_config', 'Generic parser file format configuration'],
        ['ntfl_column_mapping', 'Generic parser column-to-field mappings'],
        ['ntfl_generic_detail', 'Generic parser staging records (all vendor types)'],
    ],
    col_widths=[4.5, 10]
)

doc.add_heading('6.3 File Status IDs', level=2)
add_table(
    ['ID', 'Status', 'Description'],
    [
        ['1', 'Initial Loading', 'File is being parsed and loaded'],
        ['2', 'Transactions Loaded', 'All records successfully inserted'],
        ['3', 'Processing Errors', 'Some records failed to load'],
        ['4', 'Processing Completed', 'File fully processed'],
        ['5', 'File Discarded', 'File rejected or discarded'],
        ['10', 'File Generation In Progress', 'Output file being generated'],
        ['11', 'File Generation Complete', 'Output file ready'],
        ['12', 'Response - Some Errors', 'Response file has partial errors'],
        ['13', 'Response - No Errors', 'Response file clean'],
    ],
    col_widths=[1.5, 5, 8]
)

doc.add_heading('6.4 Transfer Status Enum', level=2)
add_table(
    ['Value', 'Name', 'Description'],
    [
        ['0', 'Pending', 'Transfer queued'],
        ['1', 'Downloading', 'File being downloaded from source'],
        ['2', 'Downloaded', 'File in Transfer folder'],
        ['3', 'Processing', 'File being parsed and loaded'],
        ['4', 'Processed', 'Successfully completed'],
        ['5', 'Error', 'Transfer or processing failed'],
        ['6', 'Skipped', 'Manually skipped by user'],
    ],
    col_widths=[1.5, 3.5, 9.5]
)

doc.add_heading('6.5 Stored Procedures', level=2)
add_table(
    ['Procedure', 'Purpose'],
    [
        ['sp_new_nt_file', 'Create a new file record, allocate file number, resolve placeholders'],
        ['ss_nt_file', 'List files with filtering (legacy 4GL query)'],
        ['sunt_file', 'Update file status'],
    ],
    col_widths=[4, 10.5]
)

# ============================================================
# 7. AUTHENTICATION & SECURITY
# ============================================================
doc.add_heading('7. Authentication & Security', level=1)

doc.add_heading('7.1 Multi-Scheme Authentication', level=2)
doc.add_paragraph(
    'The API supports two authentication methods via a "MultiAuth" policy scheme. '
    'If the request contains an X-API-Key header, the API Key handler is used; '
    'otherwise, the JWT Bearer handler is used.'
)

add_table(
    ['Method', 'Header', 'Description'],
    [
        ['JWT Bearer', 'Authorization: Bearer {token}', 'Symmetric key validation with multiple issuers, audience validation, lifetime validation'],
        ['API Key', 'X-API-Key: {key}', 'Validated via DbContextApiKeyAuthentication (database-aware, domain isolation)'],
    ],
    col_widths=[3, 5, 6.5]
)

doc.add_heading('7.2 Controller Security', level=2)
doc.add_paragraph(
    'All controllers extend DbControllerBase<FileLoaderDbContext> which provides '
    'CreateSecurityContext(operationId) to create a security context with user, domain, '
    'and endpoint information. This context is passed to all service methods for '
    'authorization and auditing.'
)

# ============================================================
# 8. CONFIGURATION
# ============================================================
doc.add_heading('8. Configuration', level=1)

doc.add_heading('8.1 Configuration Hierarchy', level=2)
config_items = [
    ('1. Shared config (required)', 'Windows: C:\\Selcomm\\configuration\\appsettings.shared.json\nLinux: /etc/selcomm/appsettings.shared.json\nOverride: SELCOMM_CONFIG_PATH environment variable'),
    ('2. Local appsettings.json', 'Module-specific overrides (optional)'),
    ('3. Environment-specific', 'appsettings.{Environment}.json (optional)'),
    ('4. Environment variables', 'Highest priority'),
]
for title, desc in config_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ': ').bold = True
    p.add_run(desc)

doc.add_heading('8.2 FileLoaderOptions (Batch Configuration)', level=2)
doc.add_paragraph(
    'Hierarchical configuration keyed by (domain, fileType). Resolution order: '
    'domain+fileType > domain+Default > Default+fileType > Default+Default > built-in defaults.'
)
add_code("""\
"FileLoaderOptions": {
  "Default": {
    "Default": { "BatchSize": 1000, "TransactionBatchSize": 1000, "UseStreamingMode": true },
    "CDR": { "BatchSize": 5000 },
    "CHG": { "BatchSize": 2000 }
  },
  "domain1": {
    "CDR": { "BatchSize": 10000 }
  }
}\
""")

doc.add_heading('8.3 Key Settings', level=2)
add_table(
    ['Setting', 'Source', 'Description'],
    [
        ['ConnectionStrings:Selcomm', 'Shared', 'Informix ODBC connection string'],
        ['JwtSettings:SecretKey', 'Shared', 'JWT signing key'],
        ['JwtSettings:Issuer', 'Shared', 'JWT issuer'],
        ['JwtSettings:Audience', 'Shared', 'JWT audience'],
        ['ApiKeySettings:ValidationEndpoint', 'Shared', 'API key validation URL'],
        ['DomainJwtSettings:{domain}:Issuer', 'Shared', 'Per-domain JWT issuers'],
    ],
    col_widths=[5.5, 2, 7]
)

doc.add_heading('8.4 Logging (Serilog)', level=2)
log_items = [
    'Console sink: structured output [HH:mm:ss LVL] Message',
    'File sink: daily rolling at logs/fileloading-{date}.log',
    'Enrichment: Application name ("FileLoading"), LogContext properties',
]
for item in log_items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 9. DEPENDENCIES
# ============================================================
doc.add_heading('9. Dependencies', level=1)

doc.add_heading('9.1 NuGet Packages', level=2)
add_table(
    ['Package', 'Version', 'Purpose'],
    [
        ['Microsoft.AspNetCore.Authentication.JwtBearer', '10.0.0', 'JWT bearer authentication'],
        ['Microsoft.IdentityModel.Tokens', '8.4.0', 'Token validation'],
        ['System.IdentityModel.Tokens.Jwt', '8.4.0', 'JWT token handling'],
        ['System.Data.Odbc', '10.0.0', 'ODBC database connectivity'],
        ['Swashbuckle.AspNetCore', '7.2.0', 'Swagger/OpenAPI generation'],
        ['Swashbuckle.AspNetCore.Annotations', '7.2.0', 'Swagger annotations'],
        ['Serilog', '4.2.0', 'Structured logging'],
        ['Serilog.AspNetCore', '8.0.3', 'ASP.NET Core integration'],
        ['Serilog.Sinks.Async', '2.1.0', 'Async log writing'],
        ['Serilog.Sinks.Console', '6.0.0', 'Console output'],
        ['Serilog.Sinks.File', '6.0.0', 'File output with rolling'],
        ['SSH.NET', '2024.1.0', 'SFTP client'],
        ['FluentFTP', '50.0.1', 'FTP client'],
        ['SharpZipLib', '1.4.2', 'GZip/Zip compression'],
        ['Cronos', '0.8.4', 'CRON expression parsing'],
        ['ClosedXML', '0.104.1', 'Excel (.xlsx) file reading for generic parser'],
    ],
    col_widths=[6.5, 2, 6]
)

doc.add_heading('9.2 Project References', level=2)
add_table(
    ['Project', 'Purpose'],
    [
        ['Selcomm.Data.Common', 'OdbcDbContext, DbControllerBase, DataResult, stored procedure execution, batch operations'],
        ['Selcomm.Authentication.Common', 'ApiKeyAuthenticationHandler, JWT/API key extensions, domain isolation'],
    ],
    col_widths=[5, 9.5]
)

# ============================================================
# 10. DEPLOYMENT
# ============================================================
doc.add_heading('10. Deployment', level=1)

doc.add_heading('10.1 Port Assignment', level=2)
doc.add_paragraph('This module is assigned port 5140 in the standard port allocation.')

doc.add_heading('10.2 Deployment Options', level=2)
add_table(
    ['Method', 'Use Case'],
    [
        ['Linux Systemd', 'Production deployment on 10.1.20.55 (LinWebProd0)'],
        ['Docker', 'Containerised deployment'],
        ['Local .NET', 'Development (dotnet run --urls=http://localhost:5140)'],
    ],
    col_widths=[4, 10.5]
)

doc.add_heading('10.3 Linux Production Configuration', level=2)
add_table(
    ['Setting', 'Value'],
    [
        ['Installation path', '/var/www/api/v4/file_loading_api/'],
        ['Service name', 'file_loading_api'],
        ['Runs as', 'weblocal:webusers'],
        ['Port', 'ASPNETCORE_URLS=http://0.0.0.0:5140'],
        ['Shared config', '/etc/selcomm/appsettings.shared.json'],
        ['Logs (journald)', '/var/log/file-loading-api/'],
        ['Logs (Serilog)', '/var/www/api/v4/file_loading_api/logs/'],
    ],
    col_widths=[4, 10.5]
)

doc.add_heading('10.4 Middleware Pipeline Order', level=2)
pipeline = [
    'Swagger (all environments)',
    'HTTPS Redirection',
    'CORS (AllowAny)',
    'Authentication',
    'Authorization',
    'Controller routing',
]
for i, item in enumerate(pipeline, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('10.5 JSON Serialisation', level=2)
json_items = [
    'Null values omitted (WhenWritingNull)',
    'Enums serialised as strings (JsonStringEnumConverter)',
    'Property naming: PascalCase (no camelCase policy)',
]
for item in json_items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 11. DESIGN PATTERNS
# ============================================================
doc.add_heading('11. Design Patterns', level=1)
add_table(
    ['Pattern', 'Where Used'],
    [
        ['Repository', 'IFileLoaderRepository / FileLoaderRepository abstracts all database access'],
        ['Strategy', 'IFileParser implementations for each file format; ITransferClient for each protocol; IFileRowReader for text vs Excel row reading'],
        ['Template Method', 'BaseFileParser defines the two-pass streaming workflow'],
        ['Factory', 'TransferClientFactory selects protocol-specific client'],
        ['Dependency Injection', 'All services, parsers, transfer clients registered in DI container'],
        ['Options Pattern', 'FileLoaderOptionsRoot with hierarchical domain/fileType resolution'],
        ['Hosted Service', 'FileTransferWorker as BackgroundService for scheduled transfers'],
        ['DataResult<T>', 'Consistent response wrapping with status code, data, and error info'],
    ],
    col_widths=[4, 10.5]
)

# ============================================================
# 12. KEY ARCHITECTURAL DECISIONS
# ============================================================
doc.add_heading('12. Key Architectural Decisions', level=1)

decisions = [
    (
        'Two-Pass Streaming over In-Memory Loading',
        'Parse files in two passes (validate, then insert) instead of loading all records into memory.',
        'A 400K-record CDR file would consume ~280MB in memory. Streaming reduces this to '
        '~10-50MB while still providing comprehensive validation before any database writes.'
    ),
    (
        'Configurable Batch Sizes per Domain/FileType',
        'Allow batch sizes and streaming mode to be configured hierarchically by domain and file type.',
        'Different vendors produce files of vastly different sizes. Telstra CDR files may have 400K+ '
        'records while charge files may have 50. One-size-fits-all batching wastes resources or risks timeouts.'
    ),
    (
        'AI-Friendly Validation Summaries',
        'Generate plain-English validation summaries alongside machine-readable error data.',
        'Enables AI agents and LLM-powered support tools to understand and explain file loading '
        'failures to users without parsing raw error codes.'
    ),
    (
        'Non-Breaking Error Collection',
        'Collect all errors during parsing rather than failing on the first error.',
        'Operators need to see the full scope of issues in a file to determine whether to fix and retry, '
        'partially process, or reject. Early termination hides the true error count.'
    ),
    (
        'Folder-Based Workflow over State Machine',
        'Use physical folder locations (Transfer, Processing, Processed, Errors, Skipped) to represent file lifecycle state.',
        'Makes the system state visible and debuggable via the filesystem. Operators can inspect, '
        'move, or recover files using standard file tools without database changes.'
    ),
]

for title, decision, rationale in decisions:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.add_run('Decision: ').bold = True
    p.add_run(decision)
    p = doc.add_paragraph()
    p.add_run('Rationale: ').bold = True
    p.add_run(rationale)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), 'architecture.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
